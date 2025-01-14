from docx import Document

documento = Document()
# para add um título
documento.add_heading('Título do documento', 0)

# para add um parágrafo
paragrafo = documento.add_paragraph('Um parágrafo simples')
paragrafo.add_run(' e super importante ').bold = True
paragrafo.add_run('do autor ')
paragrafo.add_run('Luan').italic = True


# adicionar heading(cabeçalho)
documento.add_heading('Título nível 1',level=1)
documento.add_heading('Título nível 2',level=2)
documento.add_heading('Título nível 3',level=3)
documento.add_heading('Título nível 4',level=4)

# Formatação de estilo
documento.add_paragraph('Formatação "No Spacing"',style='No Spacing')
documento.add_paragraph('Formatação "Heading1"', style='Heading 1')
documento.add_paragraph('Formatação "Heading 2"', style='Heading 2')
documento.add_paragraph('Formatação "Heading 3"', style='Heading 3')
documento.add_paragraph('Formatação "Title"', style='Title')
documento.add_paragraph('Formatação "Subtitle"', style='Subtitle')
documento.add_paragraph('Formatação "Quote"', style='Quote')
documento.add_paragraph('Formatação "Intense Quote"', style='Intense Quote')
documento.add_paragraph('Formatação "List Paragraph"', style='List Paragraph')
documento.add_paragraph('Primeiro item em uma lista com pontos', style='List Bullet')
documento.add_paragraph('primeiro item em uma lista numerada', style='List Number')
documento.save('demo.docx')